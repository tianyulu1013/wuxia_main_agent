from __future__ import annotations

import argparse
import json
import re
import sqlite3
from http import HTTPStatus
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, quote, unquote, urlparse


ROOT = Path(__file__).resolve().parents[1]
WEB_ROOT = ROOT / "web" / "card_browser"
DB_PATH = ROOT / "data" / "cards.sqlite"
UNIT_OVERRIDES_PATH = ROOT / "data" / "card_unit_overrides.json"
CARD_REVIEWS_PATH = ROOT / "data" / "card_reviews.json"
CARD_UNDERSTANDING_NOTES_PATH = ROOT / "data" / "review" / "card_understanding_notes.json"
CARD_EVALUATIONS_PATH = ROOT / "data" / "review" / "card_evaluations.json"
CARD_MAINTENANCE_TODOS_PATH = ROOT / "data" / "review" / "card_maintenance_todos.json"
CHANGE_CANDIDATES_PATH = ROOT / "data" / "change_candidates.json"
STRUCTURE_NOTES_PATH = ROOT / "data" / "card_structure_notes.json"
STATISTICS_PATH = ROOT / "data" / "review" / "card_database_statistics.json"
CARD_IMAGE_ALIASES_PATH = ROOT / "data" / "card_image_aliases.json"
SITE_DOCUMENTS_PATH = ROOT / "data" / "site_documents.json"
RELEASE_CARD_ROOT = ROOT / "data" / "release_images" / "cards"
ALL_UNITS_GROUP = "__all_units__"
CARD_IMAGE_INDEX: dict[str, Path] | None = None
CARD_IMAGE_INDEX_SIGNATURE: tuple[int, int] | None = None
SITE_DOCUMENT_PAYLOAD_CACHE: dict[str, tuple[tuple[str, float, int], dict[str, object]]] = {}


CATEGORY_LABELS = {
    "combat_characters": "战斗人物",
    "attached_characters": "附加人物",
    "items": "物品",
    "scenes": "场景",
    "titles": "称号",
    "deprecated": "废弃记录",
}


def connect() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def label_category(value: str | None) -> str:
    return CATEGORY_LABELS.get(value or "", value or "")


def row_to_result(row: sqlite3.Row) -> dict[str, object]:
    data = dict(row)
    data["category_label"] = label_category(data.get("category"))
    return apply_display_overrides(data)


def load_unit_overrides() -> dict[str, object]:
    if not UNIT_OVERRIDES_PATH.exists():
        return {}
    return json.loads(UNIT_OVERRIDES_PATH.read_text(encoding="utf-8"))


def load_json_file(path: Path, fallback: object) -> object:
    if not path.exists():
        return fallback
    return json.loads(path.read_text(encoding="utf-8"))


def load_site_document_entries() -> list[dict[str, object]]:
    data = load_json_file(SITE_DOCUMENTS_PATH, {})
    entries = data.get("documents", []) if isinstance(data, dict) else []
    normalized: list[dict[str, object]] = []
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        document_id = str(entry.get("id") or "").strip()
        relative_path = str(entry.get("path") or "").strip()
        if not document_id or not relative_path:
            continue
        path = (ROOT / relative_path).resolve()
        if ROOT.resolve() not in path.parents and path != ROOT.resolve():
            continue
        normalized.append(
            {
                "id": document_id,
                "title": str(entry.get("title") or document_id),
                "group": str(entry.get("group") or "资料"),
                "kind": str(entry.get("kind") or path.suffix.removeprefix(".") or "file"),
                "path": relative_path.replace("\\", "/"),
                "source_path": str(entry.get("source_path") or ""),
                "description": str(entry.get("description") or ""),
                "version": str(entry.get("version") or ""),
                "updated": str(entry.get("updated") or ""),
                "exists": path.exists(),
                "size": path.stat().st_size if path.exists() else 0,
            }
        )
    return normalized


def load_site_document_meta() -> dict[str, object]:
    data = load_json_file(SITE_DOCUMENTS_PATH, {})
    if not isinstance(data, dict):
        return {}
    return {
        "library_version": data.get("library_version", ""),
        "site_version": data.get("site_version", ""),
        "updated": data.get("updated", ""),
    }


def find_site_document(document_id: str) -> dict[str, object] | None:
    for entry in load_site_document_entries():
        if entry["id"] == document_id:
            return entry
    return None


def document_path(entry: dict[str, object]) -> Path:
    return (ROOT / str(entry["path"])).resolve()


def is_docx_heading(style_name: str, text: str) -> bool:
    normalized_style = style_name.lower()
    if normalized_style.startswith("heading") or normalized_style in {"title", "subtitle"}:
        return True
    return bool(
        re.match(r"^\u7b2c[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e\u96f6\u3007]+(?:\u90e8\u5206|\u7ae0|\u7bc7|\u8282)[\uff1a:]", text)
        or re.match(r"^\d+(?:\.\d+)+\s+\S+", text)
    )


def docx_heading_level(style_name: str, text: str) -> int:
    match = re.search(r"heading\s*(\d+)", style_name.lower())
    if match:
        return max(1, min(6, int(match.group(1))))
    if style_name.lower() == "title":
        return 1
    if style_name.lower() == "subtitle":
        return 2
    if text.startswith("\u7b2c\u4e00\u90e8\u5206") or text.startswith("\u7b2c\u4e8c\u90e8\u5206") or text.startswith("\u7b2c\u4e09\u90e8\u5206"):
        return 1
    if re.match(r"^\u7b2c[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e\u96f6\u3007]+\u7ae0[\uff1a:]", text):
        return 2
    if re.match(r"^\d+\.\d+\s+", text):
        return 3
    return 2


def read_docx_blocks(path: Path) -> list[dict[str, object]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("当前 Python 环境缺少 python-docx，无法读取 docx。") from exc
    document = Document(str(path))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if not text or style_name.lower().startswith("toc"):
            continue
        blocks.append(
            {
                "type": "heading" if is_docx_heading(style_name, text) else "paragraph",
                "level": docx_heading_level(style_name, text),
                "text": text,
            }
        )
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append({"type": "paragraph", "level": 0, "text": " | ".join(cells)})
    return blocks


def blocks_to_text(blocks: list[dict[str, object]]) -> str:
    return "\n\n".join(str(block.get("text") or "") for block in blocks if block.get("text"))


def blocks_to_sections(blocks: list[dict[str, object]]) -> list[dict[str, object]]:
    sections: list[dict[str, object]] = []
    current: dict[str, object] | None = None
    intro: list[str] = []
    for block in blocks:
        text = str(block.get("text") or "").strip()
        if not text:
            continue
        if block.get("type") == "heading":
            if current is not None:
                sections.append(current)
            elif intro:
                sections.append({"id": "intro", "title": "导言", "level": 1, "content": "\n\n".join(intro)})
                intro = []
            current = {
                "id": f"section-{len(sections) + 1}",
                "title": text,
                "level": int(block.get("level") or 2),
                "content": "",
            }
            continue
        if current is None:
            intro.append(text)
        else:
            existing = str(current.get("content") or "")
            current["content"] = f"{existing}\n\n{text}".strip()
    if current is not None:
        sections.append(current)
    elif intro:
        sections.append({"id": "intro", "title": "正文", "level": 1, "content": "\n\n".join(intro)})
    return sections


def read_docx_text(path: Path) -> str:
    return blocks_to_text(read_docx_blocks(path))


def read_markdown_sections(path: Path) -> list[dict[str, object]]:
    text = path.read_text(encoding="utf-8")
    blocks: list[dict[str, object]] = []
    buffer: list[str] = []
    for line in text.splitlines():
        match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if match:
            if buffer:
                blocks.append({"type": "paragraph", "level": 0, "text": "\n".join(buffer).strip()})
                buffer = []
            blocks.append({"type": "heading", "level": len(match.group(1)), "text": match.group(2).strip()})
        else:
            buffer.append(line)
    if buffer:
        blocks.append({"type": "paragraph", "level": 0, "text": "\n".join(buffer).strip()})
    return blocks_to_sections(blocks)


def read_site_document_sections(entry: dict[str, object]) -> list[dict[str, object]]:
    path = document_path(entry)
    if not path.exists():
        return []
    kind = str(entry.get("kind") or "").lower()
    suffix = path.suffix.lower()
    if kind == "docx" or suffix == ".docx":
        return blocks_to_sections(read_docx_blocks(path))
    if kind in {"markdown", "md"} or suffix == ".md":
        return read_markdown_sections(path)
    return []


def read_site_document_content(entry: dict[str, object]) -> str:
    path = document_path(entry)
    if not path.exists():
        return ""
    kind = str(entry.get("kind") or "").lower()
    suffix = path.suffix.lower()
    if kind == "document-json" or suffix == ".json":
        data = json.loads(path.read_text(encoding="utf-8"))
        return str(data.get("content") or "") if isinstance(data, dict) else ""
    if kind == "docx" or suffix == ".docx":
        return read_docx_text(path)
    return path.read_text(encoding="utf-8")


def site_document_summary(entry: dict[str, object]) -> dict[str, object]:
    path = document_path(entry)
    modified = path.stat().st_mtime if path.exists() else None
    return {**entry, "modified": modified}


def site_document_payload(entry: dict[str, object]) -> dict[str, object]:
    path = document_path(entry)
    signature = (str(path), path.stat().st_mtime if path.exists() else 0, path.stat().st_size if path.exists() else 0)
    cache_key = str(entry.get("id") or entry.get("path") or path)
    cached = SITE_DOCUMENT_PAYLOAD_CACHE.get(cache_key)
    if cached and cached[0] == signature:
        return cached[1]
    kind = str(entry.get("kind") or "").lower()
    suffix = path.suffix.lower()
    content = ""
    sections: list[dict[str, object]] = []
    if path.exists() and (kind == "document-json" or suffix == ".json"):
        data = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(data, dict):
            payload = {**site_document_summary(entry), **data}
            SITE_DOCUMENT_PAYLOAD_CACHE[cache_key] = (signature, payload)
            return payload
    if path.exists() and (kind == "docx" or suffix == ".docx"):
        blocks = read_docx_blocks(path)
        content = blocks_to_text(blocks)
        sections = blocks_to_sections(blocks)
    elif path.exists() and (kind in {"markdown", "md"} or suffix == ".md"):
        content = path.read_text(encoding="utf-8")
        sections = read_markdown_sections(path)
    elif path.exists():
        content = path.read_text(encoding="utf-8")
    payload = {
        **site_document_summary(entry),
        "content": content,
        "sections": sections,
    }
    SITE_DOCUMENT_PAYLOAD_CACHE[cache_key] = (signature, payload)
    return payload


def document_snippets(content: str, query: str, limit: int = 3, radius: int = 70) -> list[str]:
    if not query:
        return []
    snippets: list[str] = []
    lower_content = content.lower()
    lower_query = query.lower()
    start = 0
    while len(snippets) < limit:
        index = lower_content.find(lower_query, start)
        if index == -1:
            break
        left = max(0, index - radius)
        right = min(len(content), index + len(query) + radius)
        snippet = content[left:right].replace("\n", " ").strip()
        if left > 0:
            snippet = f"...{snippet}"
        if right < len(content):
            snippet = f"{snippet}..."
        snippets.append(snippet)
        start = index + len(query)
    return snippets


def search_site_documents(query: str) -> list[dict[str, object]]:
    q = query.strip()
    results: list[dict[str, object]] = []
    for entry in load_site_document_entries():
        payload = site_document_payload(entry)
        content = str(payload.get("content") or "")
        haystacks = [
            str(payload.get("title") or ""),
            str(payload.get("description") or ""),
            str(payload.get("group") or ""),
            str(payload.get("path") or ""),
            content,
        ]
        if not q or any(q.lower() in item.lower() for item in haystacks):
            results.append(
                {
                    **site_document_summary(entry),
                    "title": payload.get("title"),
                    "group": payload.get("group"),
                    "description": payload.get("description"),
                    "version": payload.get("version"),
                    "updated": payload.get("updated"),
                    "snippets": document_snippets(content, q) if q else [],
                    "match_count": content.lower().count(q.lower()) if q else 0,
                }
            )
    return results


def load_card_image_aliases() -> dict[str, object]:
    data = load_json_file(CARD_IMAGE_ALIASES_PATH, {})
    return data if isinstance(data, dict) else {}


def load_card_evaluations() -> dict[str, object]:
    data = load_json_file(CARD_EVALUATIONS_PATH, {})
    return data if isinstance(data, dict) else {}


def evaluation_methodology() -> dict[str, object]:
    methodology = load_card_evaluations().get("methodology", {})
    return methodology if isinstance(methodology, dict) else {}


def markdown_sections(text: object) -> list[dict[str, object]]:
    lines = str(text or "").replace("\r\n", "\n").splitlines()
    headings: list[tuple[int, int, str]] = []
    for index, line in enumerate(lines):
        match = re.match(r"^(#{2,4})\s+(.+?)\s*$", line)
        if match:
            headings.append((index, len(match.group(1)), match.group(2).strip()))
    sections: list[dict[str, object]] = []
    for position, (start, level, title) in enumerate(headings):
        end = len(lines)
        for next_start, next_level, _ in headings[position + 1:]:
            if next_level <= level:
                end = next_start
                break
        sections.append({"title": title, "level": level, "body": "\n".join(lines[start + 1:end]).strip()})
    return sections


def find_markdown_section(sections: list[dict[str, object]], *names: str) -> str:
    normalized_names = [re.sub(r"^[A-Z一二三四五六七八九十0-9.、\s]+", "", name).strip() for name in names]
    for section in sections:
        title = re.sub(r"^[A-Z一二三四五六七八九十0-9.、\s]+", "", str(section.get("title") or "")).strip()
        if any(name and name in title for name in normalized_names):
            return str(section.get("body") or "").strip()
    return ""


def markdown_items(text: object) -> list[str]:
    items: list[str] = []
    for raw_line in str(text or "").splitlines():
        line = raw_line.strip()
        if not line or re.match(r"^\|?\s*:?-{3,}", line):
            continue
        if line.startswith("|") and line.endswith("|"):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if cells and cells[0] not in {"问题", "能力", "项目"}:
                line = "；".join(cell for cell in cells if cell and cell != "—")
            else:
                continue
        line = re.sub(r"^(?:[-*+]\s+|\d+[.、]\s*)", "", line)
        line = re.sub(r"\*\*(.*?)\*\*", r"\1", line).replace("`", "").strip()
        if line:
            items.append(line)
    return items


def evaluation_dimension_score(full_text: str, label: str) -> int | None:
    patterns = [
        rf"(?:暂评)?{re.escape(label)}(?:能力)?\s*(?:为|：|:)?\s*(\d{{1,3}})",
        rf"{re.escape(label)}[^\n]{{0,30}}?[（(]?(\d{{1,3}})(?:分|[）)])",
    ]
    for pattern in patterns:
        match = re.search(pattern, full_text)
        if match:
            value = int(match.group(1))
            if 0 <= value <= 100:
                return value
    return None


def evaluation_section_score(section_text: str, full_text: str, label: str) -> int | None:
    direct = evaluation_dimension_score(full_text, label)
    if direct is not None:
        return direct
    for pattern in [r"(?:暂评|评分|约为|约|常态)\s*(\d{1,3})", r"(?:生存|能力)\s*(\d{1,3})", r"[（(]\s*(\d{1,3})\s*(?:，|,|分|）|\))"]:
        matches = re.findall(pattern, section_text)
        for raw in reversed(matches):
            value = int(raw)
            if 0 <= value <= 100:
                return value
    return None


def risk_payload(text: str) -> dict[str, object] | None:
    if not text:
        return None
    level = next((item for item in ["极高", "高", "中高", "中", "中低", "低"] if text.lstrip().startswith(item)), "未分级")
    return {"level": level, "summary": text}


def legacy_survival_fragment(text: str, label: str) -> str:
    lines = text.splitlines()
    for index, line in enumerate(lines):
        if label not in line or "**" not in line:
            continue
        qualifier_match = re.search(
            rf"{re.escape(label)}(?:能力)?\s*[：:]\s*([^*：:\n]+)",
            line,
        )
        qualifier = qualifier_match.group(1).strip() if qualifier_match else ""
        details: list[str] = []
        for following in lines[index + 1:]:
            if ("正面生存" in following or "侧面生存" in following) and label not in following:
                break
            cleaned = markdown_items(following)
            if cleaned:
                details.extend(cleaned)
        if qualifier or details:
            prefix = f"{qualifier}。" if qualifier else ""
            return (prefix + " ".join(details)).strip()
    matches = list(re.finditer(re.escape(label), text))
    if not matches:
        return ""
    fragment = text[matches[-1].end():]
    fragment = re.split(r"[；;\n]", fragment, maxsplit=1)[0]
    return re.sub(r"^[\s，,：:、/]+", "", fragment).strip()


def evaluation_summary(entry: dict[str, object]) -> dict[str, object]:
    structured = entry.get("summary")
    if isinstance(structured, dict):
        return structured
    full_text = str(entry.get("full_text") or "")
    sections = markdown_sections(full_text)
    core = find_markdown_section(sections, "核心定位", "核心玩法循环")
    overall = find_markdown_section(sections, "一句话总评")
    front = find_markdown_section(sections, "正面生存")
    combined_survival = find_markdown_section(sections, "初始评价", "生存能力")
    if not front and combined_survival:
        front = legacy_survival_fragment(combined_survival, "正面生存")
    if not front:
        front = str(entry.get("frontal_survival") or "")
        
    side = find_markdown_section(sections, "侧面生存")
    if not side and combined_survival:
        side = legacy_survival_fragment(combined_survival, "侧面生存")
    if not side:
        side = str(entry.get("lateral_survival") or "")
        
    pros_text = find_markdown_section(sections, "优点")
    pros = markdown_items(pros_text) if pros_text else markdown_items(entry.get("advantages"))
    
    cons_text = find_markdown_section(sections, "缺点", "缺点与死穴")
    cons = markdown_items(cons_text) if cons_text else markdown_items(entry.get("disadvantages"))
    
    questions_text = find_markdown_section(sections, "必要待校准问题", "待校准问题必要性", "待作者校准的问题")
    questions = [
        {"id": f"{entry.get('id')}_q{index}", "question": item, "status": "open"}
        for index, item in enumerate(markdown_items(questions_text), start=1)
        if item not in {"无", "无；已完成时序校正，无新增问题；-；-；-"}
    ]
    rules_risk = find_markdown_section(sections, "规则风险")
    if not rules_risk:
        rules_risk = str(entry.get("rules_risk") or "")
    digital_risk = find_markdown_section(sections, "电子化风险")
    if not digital_risk:
        digital_risk = str(entry.get("electronic_risk") or "")
    coverage = {
        "core_positioning": bool(core), "overall": bool(overall), "front_survival": bool(front),
        "side_survival": bool(side), "pros": bool(pros), "cons": bool(cons),
        "questions": bool(questions_text), "rules_risk": bool(rules_risk), "digital_risk": bool(digital_risk),
    }
    front_score = evaluation_section_score(front, full_text, "正面生存")
    side_score = evaluation_section_score(side, full_text, "侧面生存")
    return {
        "schema_version": 1,
        "source": "legacy_section_adapter",
        "coverage": "complete" if all(coverage.values()) else "partial",
        "missing_fields": [key for key, value in coverage.items() if not value],
        "core_positioning": core,
        "overall": overall,
        "survival": {
            "front": {"score": front_score, "summary": front},
            "side": {"score": side_score, "summary": side},
        },
        "pros": pros,
        "cons": cons,
        "questions": questions,
        "risks": {"rules": risk_payload(rules_risk), "digital": risk_payload(digital_risk)},
        "dimension_scores": {
            "strength": entry.get("strength_score"), "generality": entry.get("generality_score"),
            "front_survival": front_score,
            "side_survival": side_score,
            "burst": None, "control": None,
        },
    }


def normalized_evaluation_entries() -> list[dict[str, object]]:
    data = load_card_evaluations()
    entries = data.get("entries", [])
    normalized: list[dict[str, object]] = []
    if not isinstance(entries, list):
        return normalized
    with connect() as conn:
        card_rows = conn.execute(
            "SELECT id, title, normalized_title, category, author_group, source_work, source_sheet, source_row FROM cards"
        ).fetchall()
    cards_by_title: dict[str, dict[str, object]] = {}
    for row in card_rows:
        card = dict(row)
        cards_by_title.setdefault(str(card.get("title") or ""), card)
    for raw_entry in entries:
        if not isinstance(raw_entry, dict):
            continue
        entry = dict(raw_entry)
        card = cards_by_title.get(str(entry.get("card_title") or ""), {})
        entry["summary"] = evaluation_summary(entry)
        entry["card_id"] = card.get("id")
        entry["author_group"] = card.get("author_group")
        entry["source_work"] = card.get("source_work")
        entry["source_sheet"] = card.get("source_sheet")
        entry["source_row"] = card.get("source_row")
        entry["category_label"] = entry.get("category_label") or label_category(str(entry.get("category") or card.get("category") or ""))
        normalized.append(entry)
    return normalized


def evaluation_search_payload(params: dict[str, str]) -> dict[str, object]:
    q = params.get("q", "").strip().lower()
    scope = params.get("scope", "all")
    category = params.get("category", "")
    author = params.get("author", "")
    status = params.get("status", "")
    try:
        limit = min(max(int(params.get("limit", "500")), 1), 500)
    except ValueError:
        limit = 500
    entries = normalized_evaluation_entries()
    latest: dict[str, dict[str, object]] = {}
    for entry in entries:
        title = str(entry.get("card_title") or "")
        if title not in latest or int(entry.get("entry_number") or 0) > int(latest[title].get("entry_number") or 0):
            latest[title] = entry
    filtered: list[dict[str, object]] = []
    for entry in latest.values():
        if category and entry.get("category") != category:
            continue
        if author and entry.get("author_group") != author:
            continue
        if status and entry.get("status") != status:
            continue
        summary = entry.get("summary", {}) if isinstance(entry.get("summary"), dict) else {}
        survival = summary.get("survival", {}) if isinstance(summary.get("survival"), dict) else {}
        risks = summary.get("risks", {}) if isinstance(summary.get("risks"), dict) else {}
        questions = summary.get("questions", []) if isinstance(summary.get("questions"), list) else []
        fields = {
            "title": str(entry.get("card_title") or ""),
            "positioning": str(summary.get("core_positioning") or "") + "\n" + str(summary.get("overall") or ""),
            "survival": json.dumps(survival, ensure_ascii=False),
            "pros": "\n".join(str(item) for item in summary.get("pros", []) if item),
            "cons": "\n".join(str(item) for item in summary.get("cons", []) if item),
            "questions": "\n".join(str(item.get("question") or "") for item in questions if isinstance(item, dict)),
            "rules_risk": json.dumps(risks.get("rules"), ensure_ascii=False),
            "digital_risk": json.dumps(risks.get("digital"), ensure_ascii=False),
            "full_text": str(entry.get("full_text") or ""),
        }
        haystack = fields.get(scope, "\n".join(fields.values())) if scope != "all" else "\n".join(fields.values())
        if q and q not in haystack.lower():
            continue
        snippet_source = fields.get(scope) or fields["positioning"] or fields["full_text"]
        filtered.append({
            "id": entry.get("card_id"), "title": entry.get("card_title"), "category": entry.get("category"),
            "category_label": entry.get("category_label"), "author_group": entry.get("author_group"),
            "source_work": entry.get("source_work"), "source_sheet": entry.get("source_sheet"),
            "source_row": entry.get("source_row"), "snippet": re.sub(r"\s+", " ", snippet_source).strip()[:240],
            "strength_score": entry.get("strength_score"), "generality_score": entry.get("generality_score"),
            "status": entry.get("status"),
            "status_label": entry.get("status_label") or {
                "author_reviewed": "作者评估", "ai_unreviewed": "ai评估",
                "ai_draft": "ai草稿", "unreviewed": "未评估",
            }.get(str(entry.get("status") or ""), str(entry.get("status") or "未标状态")),
            "summary": summary, "full_text": entry.get("full_text"),
        })
    filtered.sort(key=lambda item: (-(int(item.get("strength_score") or -1)), str(item.get("title") or "")))
    return {"results": filtered[:limit], "reviewed_count": len(latest)}


def numeric_dimension_stats(values: list[int]) -> dict[str, object]:
    clean = sorted(value for value in values if isinstance(value, int) and 0 <= value <= 100)
    bins = {label: 0 for label in ["0-19", "20-39", "40-59", "60-69", "70-79", "80-89", "90-100"]}
    for value in clean:
        label = "0-19" if value < 20 else "20-39" if value < 40 else "40-59" if value < 60 else "60-69" if value < 70 else "70-79" if value < 80 else "80-89" if value < 90 else "90-100"
        bins[label] += 1
    median = None if not clean else clean[len(clean) // 2] if len(clean) % 2 else round((clean[len(clean)//2 - 1] + clean[len(clean)//2]) / 2, 1)
    return {"evaluated_count": len(clean), "average": round(sum(clean) / len(clean), 1) if clean else None, "median": median, "distribution": bins}


def evaluation_statistics_payload() -> dict[str, object]:
    search = evaluation_search_payload({})
    entries = search["results"] if isinstance(search.get("results"), list) else []
    dimensions = {
        "strength": {"label": "强度", **numeric_dimension_stats([item.get("strength_score") for item in entries])},
        "generality": {"label": "泛用性", **numeric_dimension_stats([item.get("generality_score") for item in entries])},
        "front_survival": {"label": "正面生存", **numeric_dimension_stats([item.get("summary", {}).get("dimension_scores", {}).get("front_survival") for item in entries])},
        "side_survival": {"label": "侧面生存", **numeric_dimension_stats([item.get("summary", {}).get("dimension_scores", {}).get("side_survival") for item in entries])},
        "burst": {"label": "爆发能力", **numeric_dimension_stats([])},
        "control": {"label": "控制能力", **numeric_dimension_stats([])},
    }
    with connect() as conn:
        total_cards = int(conn.execute("SELECT COUNT(*) FROM cards WHERE category <> 'deprecated'").fetchone()[0])
    status_counts: dict[str, int] = {}
    open_questions = 0
    for item in entries:
        label = str(item.get("status_label") or item.get("status") or "未标")
        status_counts[label] = status_counts.get(label, 0) + 1
        open_questions += sum(1 for question in item.get("summary", {}).get("questions", []) if question.get("status") == "open")
    return {
        "total_card_count": total_cards, "reviewed_card_count": len(entries),
        "unreviewed_card_count": max(0, total_cards - len(entries)), "open_question_count": open_questions,
        "status_counts": status_counts, "dimensions": dimensions,
    }


def card_evaluation_payload(title: object) -> dict[str, object]:
    card_title = str(title or "")
    data = load_card_evaluations()
    entries = data.get("entries", [])
    by_title = data.get("by_title", {})
    entry_ids = by_title.get(card_title, []) if isinstance(by_title, dict) else []
    entry_id_set = {str(item) for item in entry_ids} if isinstance(entry_ids, list) else set()
    matched = [
        entry for entry in entries
        if isinstance(entry, dict) and str(entry.get("id")) in entry_id_set
    ] if isinstance(entries, list) else []
    if not matched:
        return {"status": "unreviewed", "status_label": "未评估", "entries": []}

    if any(entry.get("status") == "author_reviewed" for entry in matched):
        status = "author_reviewed"
        status_label = "作者评估"
    elif any(entry.get("status") == "ai_unreviewed" for entry in matched):
        status = "ai_unreviewed"
        status_label = "ai评估"
    elif any(entry.get("status") == "ai_draft" for entry in matched):
        status = "ai_draft"
        status_label = "ai草稿"
    else:
        status = "unreviewed"
        status_label = "未评估"

    for entry in matched:
        entry["summary"] = evaluation_summary(entry)
        s = entry.get("status")
        if s == "author_reviewed":
            entry["status_label"] = "作者评估"
        elif s == "ai_unreviewed":
            entry["status_label"] = "ai评估"
        elif s == "ai_draft":
            entry["status_label"] = "ai草稿"

    return {
        "status": status,
        "status_label": status_label,
        "entries": matched,
    }


def load_review_layers(title: object) -> dict[str, object]:
    card_title = str(title or "")
    reviews_data = load_json_file(CARD_REVIEWS_PATH, {})
    understanding_data = load_json_file(CARD_UNDERSTANDING_NOTES_PATH, {})
    maintenance_data = load_json_file(CARD_MAINTENANCE_TODOS_PATH, {})
    candidates_data = load_json_file(CHANGE_CANDIDATES_PATH, {})
    reviews = {}
    understanding_note = {}
    maintenance_todos: list[object] = []
    if isinstance(reviews_data, dict):
        cards = reviews_data.get("cards", {})
        if isinstance(cards, dict):
            reviews = cards.get(card_title, {}) if isinstance(cards.get(card_title, {}), dict) else {}
    if isinstance(understanding_data, dict):
        notes = understanding_data.get("notes", {})
        if isinstance(notes, dict):
            note = notes.get(card_title, {})
            understanding_note = note if isinstance(note, dict) else {}
    if isinstance(maintenance_data, dict):
        todos = maintenance_data.get("todos", {})
        if isinstance(todos, dict):
            items = todos.get(card_title, [])
            maintenance_todos = items if isinstance(items, list) else []
    candidates: list[object] = []
    if isinstance(candidates_data, dict) and isinstance(candidates_data.get("candidates"), list):
        candidates = [
            item
            for item in candidates_data["candidates"]
            if isinstance(item, dict) and item.get("card_title") == card_title
        ]
    return {
        "review": reviews,
        "understanding_note": understanding_note,
        "evaluation": card_evaluation_payload(card_title),
        "maintenance_todos": maintenance_todos,
        "change_candidates": candidates,
    }


def load_structure_notes(title: object) -> list[object]:
    card_title = str(title or "")
    data = load_json_file(STRUCTURE_NOTES_PATH, {})
    if not isinstance(data, dict):
        return []
    cards = data.get("cards", {})
    if not isinstance(cards, dict):
        return []
    notes = cards.get(card_title, [])
    return notes if isinstance(notes, list) else []


def normalize_image_key(value: str) -> str:
    return re.sub(r"[\s_]+", "", value)


def load_card_image_index() -> dict[str, Path]:
    global CARD_IMAGE_INDEX, CARD_IMAGE_INDEX_SIGNATURE
    paths = list(RELEASE_CARD_ROOT.rglob("*.png")) if RELEASE_CARD_ROOT.exists() else []
    signature = (len(paths), max((path.stat().st_mtime_ns for path in paths), default=0))
    if CARD_IMAGE_INDEX is not None and CARD_IMAGE_INDEX_SIGNATURE == signature:
        return CARD_IMAGE_INDEX
    index: dict[str, Path] = {}
    for path in paths:
        name = re.sub(r"^\d+_", "", path.stem)
        candidates = {name, name.replace("_", "")}
        if "_" in name:
            candidates.update(part for part in name.split("_") if part)
        for candidate in candidates:
            index.setdefault(normalize_image_key(candidate), path)
    CARD_IMAGE_INDEX = index
    CARD_IMAGE_INDEX_SIGNATURE = signature
    return index


def image_lookup_names(card_or_title: object) -> list[str]:
    aliases = load_card_image_aliases()
    by_title = aliases.get("by_title", {}) if isinstance(aliases.get("by_title"), dict) else {}
    by_location = aliases.get("by_location", {}) if isinstance(aliases.get("by_location"), dict) else {}
    if isinstance(card_or_title, dict):
        title = str(card_or_title.get("title") or "")
        location = f"{card_or_title.get('source_sheet')}!{card_or_title.get('source_row')}"
        names = [str(by_location.get(location) or ""), str(by_title.get(title) or ""), title]
    else:
        title = str(card_or_title or "")
        names = [str(by_title.get(title) or ""), title]
    return [name for name in names if name]


def find_card_image(card_or_title: object) -> Path | None:
    keys = [normalize_image_key(name) for name in image_lookup_names(card_or_title)]
    key = next((item for item in keys if item), "")
    if not key:
        return None
    index = load_card_image_index()
    for item in keys:
        if item in index:
            return index[item]
    return None


def text_matches(value: object, needle: str) -> bool:
    if value is None:
        return False
    if isinstance(value, str):
        return needle in value
    if isinstance(value, list):
        return any(text_matches(item, needle) for item in value)
    if isinstance(value, dict):
        return any(text_matches(item, needle) for item in value.values())
    return needle in str(value)


def matching_unit_override_titles(scope: str, needle: str) -> list[str]:
    if not needle:
        return []
    matches: list[str] = []
    for title, config in load_unit_overrides().items():
        if not isinstance(config, dict):
            continue
        values: list[object] = []
        units = [unit for unit in config.get("units", []) if isinstance(unit, dict)]
        shared = config.get("shared") if isinstance(config.get("shared"), dict) else {}
        if scope == "identity":
            for unit in units:
                values.extend([unit.get("identity"), unit.get("entity_kind")])
        elif scope == "weapons":
            for unit in units:
                values.append(unit.get("weapons"))
        elif scope == "relationships":
            for unit in units:
                values.append(unit.get("relationships"))
            values.append(config.get("relationships"))
            values.append(shared.get("relationships") if isinstance(shared, dict) else None)
        elif scope == "all":
            values.append(config)
        if any(text_matches(value, needle) for value in values):
            matches.append(str(title))
    return matches


def apply_display_overrides(card: dict[str, object]) -> dict[str, object]:
    overrides = load_unit_overrides().get(str(card.get("title") or ""), {})
    if isinstance(overrides, dict) and overrides.get("suppress_card_life") and card.get("life"):
        card["source_life"] = card["life"]
        card["life"] = None
    return card


def unit_key(owner_units: object) -> str | None:
    if not isinstance(owner_units, list) or not owner_units:
        return None
    return "、".join(str(item) for item in owner_units)


def build_card_units(card: dict[str, object], abilities: list[dict[str, object]]) -> list[dict[str, object]]:
    overrides = load_unit_overrides().get(str(card.get("title") or ""), {})
    ordered: list[str] = []
    units: dict[str, dict[str, object]] = {}

    def ensure(name: str, seed: dict[str, object] | None = None) -> dict[str, object]:
        if name not in units:
            ordered.append(name)
            units[name] = {
                "name": name,
                "life": None,
                "life_pool": None,
                "gender": None,
                "entity_kind": None,
                "identity": None,
                "weapons": [],
                "relationships": None,
                "is_ability_group": False,
                "abilities": [],
            }
        if seed:
            for key in (
                "display_name",
                "life",
                "life_pool",
                "gender",
                "counts_as_characters",
                "entity_kind",
                "identity",
                "relationships",
                "is_ability_group",
                "note",
                "name_status",
            ):
                if seed.get(key):
                    units[name][key] = seed[key]
            if seed.get("weapons"):
                units[name]["weapons"] = seed["weapons"]
        return units[name]

    shared = overrides.get("shared")
    if isinstance(shared, dict) and shared:
        ensure(ALL_UNITS_GROUP, {"display_name": shared.get("display_name") or "共同特技", **shared})

    for unit in overrides.get("units", []):
        name = str(unit.get("name") or "").strip()
        if name:
            ensure(name, unit)

    explicit_unit_names = [
        str(unit.get("name") or "").strip()
        for unit in overrides.get("units", []) or []
        if isinstance(unit, dict) and str(unit.get("name") or "").strip()
    ]

    def display_unit_key(owner_units: object) -> str | None:
        if (
            len(explicit_unit_names) > 1
            and isinstance(owner_units, list)
            and set(str(item) for item in owner_units) == set(explicit_unit_names)
        ):
            return ALL_UNITS_GROUP
        return unit_key(owner_units)

    unassigned: list[dict[str, object]] = []
    for ability in abilities:
        key = display_unit_key(ability.get("owner_units"))
        if not key:
            unassigned.append(ability)
            continue
        seed = None
        owner_units = ability.get("owner_units")
        if (
            key != ALL_UNITS_GROUP
            and key not in explicit_unit_names
            and isinstance(owner_units, list)
            and len(owner_units) > 1
        ):
            seed = {"display_name": f"{key}共同特技", "is_ability_group": True}
        unit = ensure(key, seed)
        if ability.get("owner_identity") and not unit.get("identity"):
            unit["identity"] = ability["owner_identity"]
        if ability.get("owner_weapons") and not unit.get("weapons"):
            unit["weapons"] = ability["owner_weapons"]
        unit["abilities"].append(ability)

    if units and unassigned:
        ensure("未分配")["abilities"] = unassigned

    return [units[name] for name in ordered if units[name].get("abilities") or name != "未分配"]


class CardBrowserHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(WEB_ROOT), **kwargs)

    def log_message(self, format: str, *args: object) -> None:
        print(f"[card-browser] {self.address_string()} - {format % args}")

    def end_headers(self) -> None:
        self.send_header("Cache-Control", "no-store")
        super().end_headers()

    def send_json(self, payload: object, status: HTTPStatus = HTTPStatus.OK) -> None:
        body = json.dumps(payload, ensure_ascii=False, default=str).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def send_error_json(self, message: str, status: HTTPStatus = HTTPStatus.BAD_REQUEST) -> None:
        self.send_json({"error": message}, status)

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/api/meta":
            self.handle_meta()
            return
        if parsed.path == "/api/statistics":
            self.handle_statistics()
            return
        if parsed.path == "/api/documents":
            self.handle_documents()
            return
        if parsed.path == "/api/document-search":
            self.handle_document_search(parsed.query)
            return
        if parsed.path.startswith("/api/document/"):
            self.handle_document(unquote(parsed.path.removeprefix("/api/document/")))
            return
        if parsed.path == "/api/stat-query":
            self.handle_stat_query(parsed.query)
            return
        if parsed.path == "/api/evaluation-search":
            self.handle_evaluation_search(parsed.query)
            return
        if parsed.path == "/api/evaluation-stats":
            self.send_json(evaluation_statistics_payload())
            return
        if parsed.path == "/api/search":
            self.handle_search(parsed.query)
            return
        if parsed.path.startswith("/api/card-image/"):
            self.handle_card_image(unquote(parsed.path.removeprefix("/api/card-image/")))
            return
        if parsed.path.startswith("/api/card/"):
            self.handle_card(unquote(parsed.path.removeprefix("/api/card/")))
            return
        return super().do_GET()

    def handle_meta(self) -> None:
        with connect() as conn:
            metadata = {row["key"]: row["value"] for row in conn.execute("SELECT key, value FROM metadata")}
            by_category = [
                {
                    "category": row["category"],
                    "category_label": label_category(row["category"]),
                    "count": row["count"],
                }
                for row in conn.execute(
                    "SELECT category, COUNT(*) AS count FROM cards GROUP BY category ORDER BY category"
                )
            ]
            categories = [
                {"value": row["category"], "label": label_category(row["category"])}
                for row in conn.execute("SELECT DISTINCT category FROM cards ORDER BY category")
            ]
            authors = [
                row["author_group"]
                for row in conn.execute(
                    "SELECT DISTINCT author_group FROM cards WHERE author_group IS NOT NULL AND author_group <> '' ORDER BY author_group"
                )
            ]
        self.send_json(
            {
                **load_site_document_meta(),
                "evaluation_methodology": evaluation_methodology(),
                "source_workbook": metadata.get("source_workbook", ""),
                "source_path": metadata.get("source_path", ""),
                "record_count": int(metadata.get("record_count", "0")),
                "by_category": by_category,
                "categories": categories,
                "authors": authors,
            }
        )

    def handle_statistics(self) -> None:
        data = load_json_file(STATISTICS_PATH, {})
        self.send_json(data if isinstance(data, dict) else {})

    def handle_documents(self) -> None:
        self.send_json(
            {
                **load_site_document_meta(),
                "documents": [site_document_summary(entry) for entry in load_site_document_entries()],
            }
        )

    def handle_document_search(self, query_string: str) -> None:
        params = parse_qs(query_string)
        q = (params.get("q", [""])[0] or "").strip()
        self.send_json({"results": search_site_documents(q)})

    def handle_document(self, document_id: str) -> None:
        entry = find_site_document(document_id)
        if entry is None:
            self.send_error_json("未找到资料", HTTPStatus.NOT_FOUND)
            return
        try:
            self.send_json(site_document_payload(entry))
        except RuntimeError as exc:
            self.send_error_json(str(exc), HTTPStatus.INTERNAL_SERVER_ERROR)

    def search_clauses(self, params: dict[str, list[str]]) -> tuple[list[str], list[object], dict[str, object]]:
        q = (params.get("q", [""])[0] or "").strip()
        scope = (params.get("scope", ["all"])[0] or "all").strip()
        ability_type = (params.get("ability_type", [""])[0] or "").strip()
        category = (params.get("category", [""])[0] or "").strip()
        author = (params.get("author", [""])[0] or "").strip()
        clauses = []
        values: list[object] = []
        if q:
            like = f"%{q}%"
            normalized = q.replace("（", "(").replace("）", ")").replace(" ", "")
            normalized_like = f"%{normalized}%"
            override_titles = matching_unit_override_titles(scope, q)

            def append_search_clause(sql_fragment: str, fragment_values: list[object]) -> None:
                if override_titles:
                    placeholders = ", ".join("?" for _ in override_titles)
                    clauses.append(f"({sql_fragment} OR title IN ({placeholders}))")
                    values.extend(fragment_values)
                    values.extend(override_titles)
                else:
                    clauses.append(sql_fragment)
                    values.extend(fragment_values)

            if scope == "title":
                clauses.append("(title LIKE ? OR normalized_title LIKE ?)")
                values.extend([like, normalized_like])
            elif scope == "identity":
                append_search_clause(
                    """
                    (
                      identity LIKE ?
                      OR EXISTS (
                        SELECT 1 FROM card_abilities a
                        WHERE a.card_id = cards.id AND a.owner_identity LIKE ?
                      )
                    )
                    """,
                    [like, like],
                )
            elif scope == "weapons":
                append_search_clause(
                    """
                    (
                      weapons LIKE ?
                      OR EXISTS (
                        SELECT 1 FROM card_abilities a
                        WHERE a.card_id = cards.id AND a.owner_weapons_json LIKE ?
                      )
                    )
                    """,
                    [like, like],
                )
            elif scope == "source_work":
                clauses.append("source_work LIKE ?")
                values.append(like)
            elif scope == "relationships":
                append_search_clause("relationships LIKE ?", [like])
            elif scope == "ability":
                if ability_type:
                    clauses.append(
                        """
                        EXISTS (
                          SELECT 1 FROM card_abilities a
                          WHERE a.card_id = cards.id
                            AND a.kind = ?
                            AND (
                              a.name LIKE ? OR a.raw_name LIKE ?
                              OR a.type_prefix LIKE ? OR a.text LIKE ?
                            )
                        )
                        """
                    )
                    values.extend([ability_type, like, like, like, like])
                else:
                    clauses.append(
                        """
                        EXISTS (
                          SELECT 1 FROM card_abilities a
                          WHERE a.card_id = cards.id
                            AND (
                              a.kind LIKE ? OR a.name LIKE ? OR a.raw_name LIKE ?
                              OR a.type_prefix LIKE ? OR a.text LIKE ?
                            )
                        )
                        """
                    )
                    values.extend([like, like, like, like, like])
            else:
                append_search_clause(
                    """
                    (
                      title LIKE ? OR normalized_title LIKE ? OR description LIKE ? OR relationships LIKE ?
                      OR identity LIKE ? OR weapons LIKE ? OR source_work LIKE ? OR author_group LIKE ? OR all_text LIKE ?
                    )
                    """,
                    [like, normalized_like, like, like, like, like, like, like, like],
                )
        if not q and scope == "ability" and ability_type:
            clauses.append(
                """
                EXISTS (
                  SELECT 1 FROM card_abilities a
                  WHERE a.card_id = cards.id AND a.kind = ?
                )
                """
            )
            values.append(ability_type)
        if category:
            clauses.append("category = ?")
            values.append(category)
        else:
            clauses.append("category <> ?")
            values.append("deprecated")
        if author:
            clauses.append("author_group = ?")
            values.append(author)
        return clauses, values, {"q": q, "scope": scope, "ability_type": ability_type, "category": category, "author": author}

    def handle_stat_query(self, query_string: str) -> None:
        params = parse_qs(query_string)
        clauses, values, filters = self.search_clauses(params)
        where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
        with connect() as conn:
            card_rows = [dict(row) for row in conn.execute(f"SELECT id, category, author_group, source_work FROM cards {where}", values)]
            card_ids = [row["id"] for row in card_rows]
            if card_ids:
                placeholders = ", ".join("?" for _ in card_ids)
                ability_rows = [
                    dict(row)
                    for row in conn.execute(
                        f"SELECT kind, name, text FROM card_abilities WHERE card_id IN ({placeholders})",
                        card_ids,
                    )
                ]
            else:
                ability_rows = []

        def count_by(rows: list[dict[str, object]], key: str, labeler=None) -> dict[str, int]:
            counts: dict[str, int] = {}
            for row in rows:
                raw_value = str(row.get(key) or "未标")
                value = labeler(raw_value) if labeler else raw_value
                counts[value] = counts.get(value, 0) + 1
            return dict(sorted(counts.items(), key=lambda item: item[1], reverse=True))

        exclusive_count = sum(1 for row in ability_rows if "【" in str(row.get("name") or "") and "】" in str(row.get("name") or ""))
        identity_count = sum(1 for row in ability_rows if re.search(r"[（(]身份[）)]\s*$", str(row.get("text") or "").strip()))
        self.send_json(
            {
                "filters": filters,
                "card_count": len(card_rows),
                "ability_count": len(ability_rows),
                "exclusive_ability_count": exclusive_count,
                "identity_ability_count": identity_count,
                "category_counts": count_by(card_rows, "category", label_category),
                "author_counts": count_by(card_rows, "author_group"),
                "source_work_counts": count_by(card_rows, "source_work"),
                "ability_kind_counts": count_by(ability_rows, "kind"),
            }
        )

    def handle_evaluation_search(self, query_string: str) -> None:
        params = parse_qs(query_string)
        payload = evaluation_search_payload({
            "q": (params.get("q", [""])[0] or "").strip(),
            "scope": (params.get("scope", ["all"])[0] or "all").strip(),
            "category": (params.get("category", [""])[0] or "").strip(),
            "author": (params.get("author", [""])[0] or "").strip(),
            "limit": (params.get("limit", ["500"])[0] or "500").strip(),
            "status": (params.get("status", [""])[0] or "").strip(),
        })
        self.send_json(payload)

    def handle_search(self, query_string: str) -> None:
        params = parse_qs(query_string)
        q = (params.get("q", [""])[0] or "").strip()
        scope = (params.get("scope", ["all"])[0] or "all").strip()
        ability_type = (params.get("ability_type", [""])[0] or "").strip()
        category = (params.get("category", [""])[0] or "").strip()
        author = (params.get("author", [""])[0] or "").strip()
        sort = (params.get("sort", ["sheet"])[0] or "sheet").strip()
        try:
            limit = min(max(int(params.get("limit", ["60"])[0]), 1), 500)
        except ValueError:
            limit = 60

        clauses = []
        values: list[object] = []
        if q:
            like = f"%{q}%"
            normalized = q.replace("（", "(").replace("）", ")").replace(" ", "")
            normalized_like = f"%{normalized}%"
            override_titles = matching_unit_override_titles(scope, q)

            def append_search_clause(sql_fragment: str, fragment_values: list[object]) -> None:
                if override_titles:
                    placeholders = ", ".join("?" for _ in override_titles)
                    clauses.append(f"({sql_fragment} OR title IN ({placeholders}))")
                    values.extend(fragment_values)
                    values.extend(override_titles)
                else:
                    clauses.append(sql_fragment)
                    values.extend(fragment_values)

            if scope == "title":
                clauses.append("(title LIKE ? OR normalized_title LIKE ?)")
                values.extend([like, normalized_like])
            elif scope == "identity":
                append_search_clause(
                    """
                    (
                      identity LIKE ?
                      OR EXISTS (
                        SELECT 1 FROM card_abilities a
                        WHERE a.card_id = cards.id AND a.owner_identity LIKE ?
                      )
                    )
                    """,
                    [like, like],
                )
            elif scope == "weapons":
                append_search_clause(
                    """
                    (
                      weapons LIKE ?
                      OR EXISTS (
                        SELECT 1 FROM card_abilities a
                        WHERE a.card_id = cards.id AND a.owner_weapons_json LIKE ?
                      )
                    )
                    """,
                    [like, like],
                )
            elif scope == "source_work":
                clauses.append("source_work LIKE ?")
                values.append(like)
            elif scope == "relationships":
                append_search_clause("relationships LIKE ?", [like])
            elif scope == "ability":
                if ability_type:
                    clauses.append(
                        """
                        EXISTS (
                          SELECT 1 FROM card_abilities a
                          WHERE a.card_id = cards.id
                            AND a.kind = ?
                            AND (
                              a.name LIKE ? OR a.raw_name LIKE ?
                              OR a.type_prefix LIKE ? OR a.text LIKE ?
                            )
                        )
                        """
                    )
                    values.extend([ability_type, like, like, like, like])
                else:
                    clauses.append(
                        """
                        EXISTS (
                          SELECT 1 FROM card_abilities a
                          WHERE a.card_id = cards.id
                            AND (
                              a.kind LIKE ? OR a.name LIKE ? OR a.raw_name LIKE ?
                              OR a.type_prefix LIKE ? OR a.text LIKE ?
                            )
                        )
                        """
                    )
                    values.extend([like, like, like, like, like])
            else:
                append_search_clause(
                    """
                    (
                      title LIKE ? OR normalized_title LIKE ? OR description LIKE ? OR relationships LIKE ?
                      OR identity LIKE ? OR weapons LIKE ? OR source_work LIKE ? OR author_group LIKE ? OR all_text LIKE ?
                    )
                    """,
                    [like, normalized_like, like, like, like, like, like, like, like],
                )
        if not q and scope == "ability" and ability_type:
            clauses.append(
                """
                EXISTS (
                  SELECT 1 FROM card_abilities a
                  WHERE a.card_id = cards.id AND a.kind = ?
                )
                """
            )
            values.append(ability_type)
        if category:
            clauses.append("category = ?")
            values.append(category)
        else:
            clauses.append("category <> ?")
            values.append("deprecated")
        if author:
            clauses.append("author_group = ?")
            values.append(author)

        where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
        base_order = {
            "title": "title COLLATE NOCASE, source_sheet, source_row",
            "category": "category, source_sheet, source_row",
            "sheet": "source_sheet, source_row",
        }.get(sort, "source_sheet, source_row")
        if q:
            order_by = (
                "CASE "
                "WHEN title = ? THEN 0 "
                "WHEN normalized_title = ? THEN 1 "
                "WHEN title LIKE ? THEN 2 "
                "ELSE 3 END, "
                f"{base_order}"
            )
            values.extend([q, normalized, f"%{q}%"])
        else:
            order_by = base_order
        values.append(limit)

        sql = f"""
            SELECT id, title, category, source_sheet, source_row, author_group, source_work,
                   life, identity, weapons, description, relationships,
                   CASE
                     WHEN description IS NOT NULL AND description <> '' THEN description
                     WHEN relationships IS NOT NULL AND relationships <> '' THEN relationships
                     ELSE all_text
                   END AS snippet
            FROM cards
            {where}
            ORDER BY {order_by}
            LIMIT ?
        """
        with connect() as conn:
            rows = [row_to_result(row) for row in conn.execute(sql, values)]
        if scope in {"identity", "weapons", "source_work", "relationships"}:
            snippet_field = {
                "identity": "identity",
                "weapons": "weapons",
                "source_work": "source_work",
                "relationships": "relationships",
            }[scope]
            for row in rows:
                row["snippet"] = row.get(snippet_field) or row.get("snippet")
        self.send_json({"results": rows})

    def handle_card(self, card_id: str) -> None:
        with connect() as conn:
            row = conn.execute("SELECT * FROM cards WHERE id = ?", (card_id,)).fetchone()
            ability_rows = conn.execute(
                """
                SELECT id, ordinal, kind, name, raw_name, type_prefix, source_field,
                       start_line, end_line, text, is_exclusive, is_identity,
                       owner_units_json, owner_identity, owner_weapons_json, review_flags_json
                FROM card_abilities
                WHERE card_id = ?
                ORDER BY ordinal
                """,
                (card_id,),
            ).fetchall()
        if row is None:
            self.send_error_json("未找到卡牌", HTTPStatus.NOT_FOUND)
            return
        payload = row_to_result(row)
        payload["abilities"] = [
            {
                **dict(ability),
                "is_exclusive": bool(ability["is_exclusive"]),
                "is_identity": bool(ability["is_identity"]),
                "owner_units": json.loads(ability["owner_units_json"]) if ability["owner_units_json"] else None,
                "owner_identity": ability["owner_identity"],
                "owner_weapons": json.loads(ability["owner_weapons_json"]) if ability["owner_weapons_json"] else None,
                "review_flags": json.loads(ability["review_flags_json"]),
            }
            for ability in ability_rows
        ]
        payload["units"] = build_card_units(payload, payload["abilities"])
        payload.update(load_review_layers(payload.get("title")))
        payload["structure_notes"] = load_structure_notes(payload.get("title"))
        if find_card_image(payload):
            payload["image_url"] = f"/api/card-image/{quote(str(payload['id']))}"
        self.send_json(payload)

    def handle_card_image(self, card_id: str) -> None:
        with connect() as conn:
            row = conn.execute("SELECT * FROM cards WHERE id = ?", (card_id,)).fetchone()
        if row is None:
            self.send_error_json("未找到卡牌", HTTPStatus.NOT_FOUND)
            return
        card = row_to_result(row)
        path = find_card_image(card)
        if path is None or not path.exists():
            self.send_error_json("未找到卡面图片", HTTPStatus.NOT_FOUND)
            return
        body = path.read_bytes()
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", "image/png")
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Cache-Control", "no-cache")
        self.end_headers()
        self.wfile.write(body)


def main() -> None:
    parser = argparse.ArgumentParser(description="Serve the local card browser.")
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", type=int, default=8765)
    args = parser.parse_args()

    if not DB_PATH.exists():
        raise FileNotFoundError(f"Database not found: {DB_PATH}")
    server = ThreadingHTTPServer((args.host, args.port), CardBrowserHandler)
    print(f"Card browser: http://{args.host}:{args.port}")
    server.serve_forever()


if __name__ == "__main__":
    main()
